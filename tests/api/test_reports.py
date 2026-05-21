"""Integration tests for /api/reports/ endpoints."""

from io import BytesIO
import uuid

import httpx
import pytest
from docx import Document
from openpyxl import load_workbook

pytestmark = [pytest.mark.asyncio, pytest.mark.api]

BASE = "/api/reports/"


def _assert_attachment(download: httpx.Response, filename: str, content_type_fragment: str) -> None:
    assert download.status_code == 200, download.text
    assert content_type_fragment in download.headers.get("content-type", "")
    assert filename in download.headers.get("content-disposition", "")
    assert download.content, f"{filename} downloaded as an empty file"


def _report_headers() -> dict[str, str]:
    octet_a = int(uuid.uuid4().hex[:2], 16) % 254 + 1
    octet_b = int(uuid.uuid4().hex[2:4], 16) % 254 + 1
    return {"X-Forwarded-For": f"10.250.{octet_a}.{octet_b}"}


async def _create_template(admin_client: httpx.AsyncClient) -> str:
    resp = await admin_client.post(
        "/api/test-templates/",
        json={
            "name": f"report-api-{uuid.uuid4().hex[:6]}",
            "description": "Minimal template for report smoke tests",
            "test_ids": ["U01"],
            "device_category": "camera",
        },
    )
    assert resp.status_code == 201, resp.text
    return resp.json()["id"]


async def _create_completed_run(admin_client: httpx.AsyncClient, device_id: str, template_id: str) -> str:
    run_resp = await admin_client.post(
        "/api/test-runs/",
        json={
            "device_id": device_id,
            "template_id": template_id,
            "connection_scenario": "direct",
        },
    )
    assert run_resp.status_code == 201, run_resp.text
    run_id = run_resp.json()["id"]

    results_resp = await admin_client.get("/api/test-results/", params={"test_run_id": run_id})
    assert results_resp.status_code == 200, results_resp.text
    results = results_resp.json()
    assert results, "Expected generated test results for the new run"

    for result in results:
        update_resp = await admin_client.post(
            f"/api/test-results/{result['id']}/override",
            json={
                "verdict": "pass",
                "comment": "Integration report smoke test",
                "override_reason": "Integration smoke test completed automatically",
            },
        )
        assert update_resp.status_code == 200, update_resp.text

    complete_resp = await admin_client.post(f"/api/test-runs/{run_id}/complete")
    assert complete_resp.status_code == 200, complete_resp.text
    return run_id


async def test_generate_generic_excel_report(admin_client: httpx.AsyncClient, test_device: dict):
    template_id = await _create_template(admin_client)
    try:
        run_id = await _create_completed_run(admin_client, test_device["id"], template_id)

        resp = await admin_client.post(
            f"{BASE}generate",
            json={
                "test_run_id": run_id,
                "report_type": "excel",
                "template_key": "generic",
            },
            headers=_report_headers(),
        )
        assert resp.status_code == 200, resp.text
        body = resp.json()
        assert body["report_type"] == "excel"
        assert body["template_key"] == "generic"
        assert body["filename"].endswith(".xlsx")

        download = await admin_client.get(f"{BASE}download/{body['filename']}")
        _assert_attachment(download, body["filename"], "spreadsheetml")
        assert download.content.startswith(b"PK\x03\x04")

        workbook = load_workbook(BytesIO(download.content))
        try:
            assert workbook.sheetnames == [
                "TEST SUMMARY",
                "TESTPLAN",
                "ADDITIONAL INFORMATION",
            ]
            assert workbook["TESTPLAN"]["B4"].value == "U01"
            assert workbook["TESTPLAN"]["F4"].value == "PASS"
            assert "Integration report smoke test" in workbook["TESTPLAN"]["G4"].value
            assert workbook["TESTPLAN"]["L3"].value == "Raw Evidence"
            assert "Raw Evidence" not in workbook.sheetnames
        finally:
            workbook.close()
    finally:
        await admin_client.delete(f"/api/test-templates/{template_id}")


@pytest.mark.parametrize("report_type", ["pdf", "csv", "cad", "dxf"])
async def test_generate_report_rejects_removed_formats(admin_client: httpx.AsyncClient, report_type: str):
    resp = await admin_client.post(
        f"{BASE}generate",
        json={
            "test_run_id": str(uuid.uuid4()),
            "report_type": report_type,
            "template_key": "generic",
        },
        headers=_report_headers(),
    )

    assert resp.status_code == 422, resp.text


async def test_generate_word_report_returns_template_key(admin_client: httpx.AsyncClient, test_device: dict):
    template_id = await _create_template(admin_client)
    try:
        run_id = await _create_completed_run(admin_client, test_device["id"], template_id)

        resp = await admin_client.post(
            f"{BASE}generate",
            json={
                "test_run_id": run_id,
                "report_type": "word",
                "template_key": "generic",
            },
            headers=_report_headers(),
        )
        assert resp.status_code == 200, resp.text
        body = resp.json()
        assert body["report_type"] == "word"
        assert body["template_key"] == "generic"
        assert body["filename"].endswith(".docx")

        download = await admin_client.get(f"{BASE}download/{body['filename']}")
        _assert_attachment(download, body["filename"], "wordprocessingml")
        assert download.content.startswith(b"PK\x03\x04")
        document = Document(BytesIO(download.content))
        full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        table_text = "\n".join(
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )
        assert "IP Device Qualification Report" in full_text
        assert "Generic IP Device" in full_text
        assert "Raw Evidence" in full_text
        assert "Detailed Evidence" in table_text
    finally:
        await admin_client.delete(f"/api/test-templates/{template_id}")


async def test_generate_report_rejects_pending_manual_results(admin_client: httpx.AsyncClient, test_device: dict):
    resp = await admin_client.post(
        "/api/test-templates/",
        json={
            "name": f"report-manual-{uuid.uuid4().hex[:6]}",
            "description": "Manual-only template for pending manual report rejection",
            "test_ids": ["U20"],
            "device_category": "camera",
        },
    )
    assert resp.status_code == 201, resp.text
    template_id = resp.json()["id"]
    try:
        run_resp = await admin_client.post(
            "/api/test-runs/",
            json={
                "device_id": test_device["id"],
                "template_id": template_id,
                "connection_scenario": "test_lab",
            },
        )
        assert run_resp.status_code == 201, run_resp.text
        run_id = run_resp.json()["id"]

        resp = await admin_client.post(
            f"{BASE}generate",
            json={
                "test_run_id": run_id,
                "report_type": "excel",
                "template_key": "generic",
            },
            headers=_report_headers(),
        )

        assert resp.status_code == 409, resp.text
        assert "manual tests are still pending" in resp.json()["detail"].lower()
    finally:
        await admin_client.delete(f"/api/test-templates/{template_id}")


async def test_generate_report_rejects_awaiting_review_runs(admin_client: httpx.AsyncClient, test_device: dict):
    template_id = await _create_template(admin_client)
    try:
        run_id = await _create_completed_run(admin_client, test_device["id"], template_id)

        review_resp = await admin_client.post(f"/api/test-runs/{run_id}/request-review")
        assert review_resp.status_code == 200, review_resp.text

        resp = await admin_client.post(
            f"{BASE}generate",
            json={
                "test_run_id": run_id,
                "report_type": "excel",
                "template_key": "generic",
            },
            headers=_report_headers(),
        )

        assert resp.status_code == 409, resp.text
        assert "reviewer sign-off" in resp.json()["detail"].lower()
    finally:
        await admin_client.delete(f"/api/test-templates/{template_id}")


async def test_generate_report_invalid_run(admin_client: httpx.AsyncClient):
    """POST /api/reports/generate with a fake test_run_id returns 404."""
    fake_id = str(uuid.uuid4())
    resp = await admin_client.post(
        f"{BASE}generate",
        json={"test_run_id": fake_id},
        headers=_report_headers(),
    )
    assert resp.status_code == 404, (
        f"Expected 404 for fake test_run_id, got {resp.status_code}: {resp.text}"
    )


async def test_report_templates(admin_client: httpx.AsyncClient):
    """GET /api/reports/templates returns available report templates."""
    resp = await admin_client.get(f"{BASE}templates")
    assert resp.status_code == 200
    body = resp.json()
    assert isinstance(body, list)
    template_keys = {item["key"] for item in body}
    assert template_keys == {"generic"}
    generic = next(item for item in body if item["key"] == "generic")
    assert generic["template_exists"] is True
    assert generic["mapping_exists"] is True
