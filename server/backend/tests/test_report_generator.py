from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
import importlib.util

import pytest
from openpyxl import load_workbook
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker

from app import main as main_module
from app.models import database as database_module
from app.models.protocol_observer_settings import ProtocolObserverSettings
from app.services import test_engine as test_engine_module
from app.services.evaluation import evaluate_result
from app.services import report_generator
from app.services.test_engine import TestEngine


def _make_test_run():
    device = SimpleNamespace(
        category="camera",
        manufacturer="Axis",
        model="P3245",
        firmware_version="10.12",
        serial_number="SN-001",
    )
    engineer = SimpleNamespace(full_name="Engineer One")
    return SimpleNamespace(
        id="12345678-1234-1234-1234-123456789abc",
        created_at=datetime(2026, 4, 10, tzinfo=timezone.utc),
        completed_at=datetime(2026, 4, 10, tzinfo=timezone.utc),
        overall_verdict="pass",
        connection_scenario="direct",
        device=device,
        engineer=engineer,
        synopsis=None,
    )


def _make_test_result():
    return SimpleNamespace(
        test_id="U01",
        test_name="Network Reachability",
        tier="automatic",
        tool="nmap",
        verdict="pass",
        comment="Device responded as expected.",
        comment_override=None,
        engineer_notes=None,
        override_reason=None,
        raw_output=None,
        parsed_data=None,
        findings=None,
        is_essential="YES",
    )


def _make_branding():
    return SimpleNamespace(
        company_name="Verdent QA",
        primary_color="#123456",
        footer_text="Confidential Qualification Report",
        logo_path="",
    )


DOCX_AVAILABLE = importlib.util.find_spec("docx") is not None
FPDF_AVAILABLE = importlib.util.find_spec("fpdf") is not None


def test_generic_template_is_available():
    templates = {item["key"]: item for item in report_generator.get_available_templates()}
    assert templates["generic"]["template_exists"] is True
    assert templates["generic"]["mapping_exists"] is True
    assert templates["sauter_680_as"]["template_exists"] is True
    assert templates["sauter_680_as"]["mapping_exists"] is True


def test_build_report_document_uses_template_mapping_for_layout():
    report = report_generator.build_report_document(
        _make_test_run(),
        [_make_test_result()],
        template_key="pelco_camera",
    )

    assert report.summary_section_title == "TEST SYNOPSIS"
    assert report.testplan_section_title == "TESTPLAN"
    assert report.additional_section_title == "ADDITIONAL INFO"
    assert report.summary_text_label == "Synopsis"
    assert [key for key, _ in report.summary_fields] == [
        "test_attempt",
        "date_range",
        "system",
        "manufacturer",
        "model",
        "firmware",
        "serial",
        "tester_name",
        "overall_result",
    ]
    assert [label for _, label in report.testplan_columns] == [
        "Test Number",
        "Brief Description",
        "Test Description",
        "Essential Pass",
        "Test Result",
        "Test Comments",
        "Engineer Notes",
        "Script Flag",
    ]


@pytest.mark.asyncio
async def test_generate_excel_report_preserves_generic_workbook_structure(tmp_path, monkeypatch):
    monkeypatch.setattr(report_generator.settings, "REPORT_DIR", str(tmp_path))

    output = await report_generator.generate_excel_report(
        _make_test_run(),
        [_make_test_result()],
        template_key="generic",
    )

    workbook = load_workbook(output)
    try:
        assert workbook.sheetnames == [
            "General Test Information",
            "Test Results",
            "Additional Device Information",
            "Raw Evidence",
        ]
        assert len(workbook["General Test Information"]._images) >= 1
    finally:
        workbook.close()


@pytest.mark.asyncio
async def test_generate_excel_report_patches_sauter_template(tmp_path, monkeypatch):
    monkeypatch.setattr(report_generator.settings, "REPORT_DIR", str(tmp_path))

    output = await report_generator.generate_excel_report(
        _make_test_run(),
        [_make_test_result()],
        template_key="sauter_680_as",
    )

    workbook = load_workbook(output)
    try:
        assert workbook.sheetnames == ["TEST SUMMARY", "TESTPLAN"]
        summary_ws = workbook["TEST SUMMARY"]
        plan_ws = workbook["TESTPLAN"]
        assert summary_ws["G11"].value == "Axis"
        assert summary_ws["G16"].value == "PASS"
        assert plan_ws["F17"].value == "PASS"
        assert plan_ws["G17"].value == "Device responded as expected."
        assert len(summary_ws._images) >= 1
    finally:
        workbook.close()


@pytest.mark.asyncio
async def test_generate_excel_report_adds_logo_to_template_summary_sheet(tmp_path, monkeypatch):
    monkeypatch.setattr(report_generator.settings, "REPORT_DIR", str(tmp_path))

    output = await report_generator.generate_excel_report(
        _make_test_run(),
        [_make_test_result()],
        template_key="pelco_camera",
    )

    workbook = load_workbook(output)
    try:
        assert len(workbook["TEST SYNOPSIS"]._images) >= 1
    finally:
        workbook.close()


def test_output_path_uses_full_run_uuid(tmp_path, monkeypatch):
    monkeypatch.setattr(report_generator.settings, "REPORT_DIR", str(tmp_path))
    run = _make_test_run()
    run.id = "12345678-1234-1234-1234-123456789abc"

    path = report_generator._output_path(run, "generic", ".xlsx")

    assert "12345678-1234-1234-1234-123456789abc" in path.name


def test_output_path_rejects_non_uuid_run_id(tmp_path, monkeypatch):
    monkeypatch.setattr(report_generator.settings, "REPORT_DIR", str(tmp_path))
    run = _make_test_run()
    run.id = "../not-a-run-id"

    with pytest.raises(ValueError):
        report_generator._output_path(run, "generic", ".xlsx")


@pytest.mark.asyncio
@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed in local test environment")
async def test_generate_word_report_includes_template_profile_and_branding(tmp_path, monkeypatch):
    from docx import Document

    monkeypatch.setattr(report_generator.settings, "REPORT_DIR", str(tmp_path))

    output = await report_generator.generate_word_report(
        _make_test_run(),
        [_make_test_result()],
        template_key="pelco_camera",
        branding_settings=_make_branding(),
    )

    document = Document(output)
    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)

    assert "Verdent QA" in full_text
    assert "Template Profile: Pelco Camera (Rev 2)" in full_text
    assert "IP Device Qualification Report" in full_text
    assert "TEST SYNOPSIS" in full_text
    assert "ADDITIONAL INFO" in full_text
    assert "Essential Pass" in table_text
    assert "Script Flag" in table_text


@pytest.mark.asyncio
@pytest.mark.skipif(not FPDF_AVAILABLE, reason="fpdf2 not installed in local test environment")
async def test_generate_pdf_report_includes_template_profile_and_branding(tmp_path, monkeypatch):
    monkeypatch.setattr(report_generator.settings, "REPORT_DIR", str(tmp_path))
    monkeypatch.setenv("EDQ_USE_LIBREOFFICE", "0")

    output = await report_generator.generate_pdf_report(
        _make_test_run(),
        [_make_test_result()],
        template_key="pelco_camera",
        branding_settings=_make_branding(),
    )

    pdf_bytes = BytesIO(Path(output).read_bytes())
    content = pdf_bytes.getvalue().decode("latin-1", errors="ignore")

    assert "IP DEVICE QUALIFICATION REPORT" in content
    assert "Pelco Camera" in content and "Rev 2" in content
    assert "TEST SYNOPSIS" in content
    assert "ADDITIONAL INFO" in content
    assert "Essential Pass" in content
    assert "Script Flag" in content


@pytest.mark.asyncio
async def test_generate_csv_report_uses_template_profile_metadata(tmp_path, monkeypatch):
    monkeypatch.setattr(report_generator.settings, "REPORT_DIR", str(tmp_path))

    output = await report_generator.generate_csv_report(
        _make_test_run(),
        [_make_test_result()],
        template_key="pelco_camera",
        branding_settings=_make_branding(),
    )

    csv_text = Path(output).read_text(encoding="utf-8")

    assert "TEST SYNOPSIS" in csv_text
    assert "Template Profile" in csv_text
    assert "Device responded as expected." in csv_text
    assert "Readiness Status" in csv_text
    assert "TESTPLAN" in csv_text
    assert "ADDITIONAL INFO" in csv_text
    assert "Synopsis" in csv_text
    assert "Essential Pass" in csv_text
    assert "Script Flag" in csv_text
    assert "Report Logo" in csv_text
    assert "electracom-logo.png" in csv_text


@pytest.mark.asyncio
async def test_generate_csv_report_keeps_client_logo_alongside_electracom_logo(tmp_path, monkeypatch):
    upload_root = tmp_path / "uploads"
    branding_dir = upload_root / "branding"
    branding_dir.mkdir(parents=True)
    source_logo = report_generator._resolve_electracom_report_logo_path()
    assert source_logo is not None
    (branding_dir / "client-logo.png").write_bytes(source_logo.read_bytes())

    monkeypatch.setattr(report_generator.settings, "REPORT_DIR", str(tmp_path / "reports"))
    monkeypatch.setattr(report_generator.settings, "UPLOAD_DIR", str(upload_root))
    branding = _make_branding()
    branding.logo_path = "client-logo.png"

    output = await report_generator.generate_csv_report(
        _make_test_run(),
        [_make_test_result()],
        template_key="generic",
        branding_settings=branding,
    )

    csv_text = Path(output).read_text(encoding="utf-8")
    assert "Report Logo" in csv_text
    assert "electracom-logo.png" in csv_text
    assert "Client Logo" in csv_text
    assert "client-logo.png" in csv_text


def test_uploaded_logo_resolution_stays_inside_upload_dir(tmp_path, monkeypatch):
    upload_root = tmp_path / "uploads"
    branding_dir = upload_root / "branding"
    branding_dir.mkdir(parents=True)
    outside_logo = tmp_path / "outside-logo.png"
    source_logo = report_generator._resolve_electracom_report_logo_path()
    assert source_logo is not None
    outside_logo.write_bytes(source_logo.read_bytes())

    monkeypatch.setattr(report_generator.settings, "UPLOAD_DIR", str(upload_root))

    assert report_generator._resolve_uploaded_logo_path("../outside-logo.png") is None


def _make_test_result_with_notes(notes: str = "Check firmware 1.2.3 manually on next reboot"):
    return SimpleNamespace(
        test_id="U01",
        test_name="Network Reachability",
        tier="guided_manual",
        tool="manual",
        verdict="pass",
        comment="Device responded as expected.",
        comment_override=None,
        engineer_notes=notes,
        override_reason=None,
        raw_output=None,
        parsed_data=None,
        findings=None,
        is_essential="YES",
    )


def _make_generic_results():
    return [
        SimpleNamespace(
            test_id="U01",
            test_name="Network Reachability",
            tier="automatic",
            tool="nmap",
            verdict="pass",
            comment="Device responded as expected.",
            comment_override=None,
            engineer_notes=None,
            override_reason=None,
            raw_output="Host is up.\n1 open tcp port discovered.",
            parsed_data=None,
            findings=None,
            is_essential="YES",
        ),
        SimpleNamespace(
            test_id="U03",
            test_name="Switch Negotiation",
            tier="guided_manual",
            tool=None,
            verdict="pass",
            comment="Engineer verified full duplex support.",
            comment_override=None,
            engineer_notes="Observed from managed switch port status page.",
            override_reason=None,
            raw_output=None,
            parsed_data=None,
            findings=None,
            is_essential="NO",
        ),
    ]


@pytest.mark.asyncio
async def test_engineer_notes_exported_as_dedicated_csv_column(tmp_path, monkeypatch):
    monkeypatch.setattr(report_generator.settings, "REPORT_DIR", str(tmp_path))
    note = "Check firmware 1.2.3 manually on next reboot"

    output = await report_generator.generate_csv_report(
        _make_test_run(),
        [_make_test_result_with_notes(note)],
        template_key="pelco_camera",
    )

    csv_text = Path(output).read_text(encoding="utf-8")

    assert "Engineer Notes" in csv_text
    assert note in csv_text
    assert "[Engineer Notes:" not in csv_text


def test_engineer_notes_resolved_as_separate_column_in_document():
    report = report_generator.build_report_document(
        _make_test_run(),
        [_make_test_result_with_notes("Inspect after reboot")],
        template_key="pelco_camera",
    )

    attrs = [attr for attr, _ in report.testplan_columns]
    assert "engineer_notes" in attrs
    row = next(item for item in report.rows if item.engineer_notes)
    assert row.engineer_notes == "Inspect after reboot"
    assert "[Engineer Notes:" not in row.test_comments


@pytest.mark.asyncio
async def test_generic_report_uses_one_row_per_test_and_keeps_evidence(tmp_path, monkeypatch):
    monkeypatch.setattr(report_generator.settings, "REPORT_DIR", str(tmp_path))

    output = await report_generator.generate_excel_report(
        _make_test_run(),
        _make_generic_results(),
        template_key="generic",
    )

    wb = load_workbook(output)
    try:
        ws = wb["Test Results"]
        headers = [ws[f"{col}3"].value for col in ["B", "C", "D", "E", "F", "G", "H", "I", "J"]]
        assert headers == [
            "Test ID",
            "Test Name",
            "Tier",
            "Tool",
            "Essential Test",
            "Test Result",
            "Test Comments",
            "Engineer Notes",
            "Evidence Summary",
        ]
        assert ws["B4"].value == "U01"
        assert ws["C4"].value == "Network Reachability"
        assert ws["J4"].value.startswith("Host is up.")
        assert ws["B5"].value == "U03"
        assert ws["D5"].value == "Guided Manual"
        assert ws["I5"].value == "Observed from managed switch port status page."
        evidence_ws = wb["Raw Evidence"]
        assert evidence_ws["A4"].value == "U01"
        assert evidence_ws["D4"].value == "Host is up.\n1 open tcp port discovered."
        assert evidence_ws["C5"].value == "Observed from managed switch port status page."
    finally:
        wb.close()


@pytest.mark.asyncio
async def test_protocol_settings_reload_drive_u04_exported_evidence(
    tmp_path,
    monkeypatch,
    db_engine,
    db_session: AsyncSession,
):
    monkeypatch.setattr(report_generator.settings, "REPORT_DIR", str(tmp_path))
    session_factory = async_sessionmaker(db_engine, class_=AsyncSession, expire_on_commit=False)
    monkeypatch.setattr(database_module, "async_session", session_factory)
    monkeypatch.setattr(report_generator.settings, "PROTOCOL_OBSERVER_ENABLED", False)
    monkeypatch.setattr(report_generator.settings, "PROTOCOL_OBSERVER_DHCP_OFFER_IP", "")

    db_session.add(
        ProtocolObserverSettings(
            singleton_key="_",
            enabled=True,
            bind_host="127.0.0.1",
            timeout_seconds=20,
            dns_port=5300,
            ntp_port=1123,
            dhcp_port=1067,
            dhcp_offer_ip="192.168.4.68",
            dhcp_subnet_mask="255.255.255.0",
            dhcp_router_ip="192.168.4.1",
            dhcp_dns_server="192.168.4.1",
            dhcp_lease_seconds=600,
        )
    )
    await db_session.commit()

    await main_module._load_protocol_observer_settings_from_db()

    assert report_generator.settings.PROTOCOL_OBSERVER_ENABLED is True
    assert report_generator.settings.PROTOCOL_OBSERVER_DHCP_OFFER_IP == "192.168.4.68"

    async def fake_observe_dhcp_activity(*, expected_mac: str | None, timeout_seconds=None, port=None):
        assert expected_mac == "AA:BB:CC:DD:EE:FF"
        assert test_engine_module.settings.PROTOCOL_OBSERVER_DHCP_OFFER_IP == "192.168.4.68"
        return {
            "observed": True,
            "lease_acknowledged": True,
            "offer_capable": True,
            "offered_ip": "192.168.4.68",
            "server_identifier": "192.168.4.1",
            "events": [
                {"message_type": 1, "observer_reply_type": 2},
                {"message_type": 3, "observer_reply_type": 5},
            ],
        }

    monkeypatch.setattr(test_engine_module, "observe_dhcp_activity", fake_observe_dhcp_activity)

    engine = TestEngine()
    parsed, raw = await engine._dispatch_test(
        "U04",
        "192.168.4.68",
        "run-u04-export",
        SimpleNamespace(mac_address="AA:BB:CC:DD:EE:FF"),
        "direct",
    )
    verdict, comment = evaluate_result("U04", parsed)

    result = SimpleNamespace(
        test_id="U04",
        test_name="DHCP Behaviour",
        tier="automatic",
        tool="protocol_observer",
        verdict=verdict,
        comment=comment,
        comment_override=None,
        engineer_notes=None,
        override_reason=None,
        raw_output=raw,
        parsed_data=parsed,
        findings=None,
        is_essential="YES",
    )

    output = await report_generator.generate_excel_report(
        _make_test_run(),
        [result],
        template_key="generic",
    )

    wb = load_workbook(output)
    try:
        results_ws = wb["Test Results"]
        evidence_ws = wb["Raw Evidence"]
        assert results_ws["B4"].value == "U04"
        assert "lease acknowledgement" in str(results_ws["H4"].value).lower()
        assert "192.168.4.68" in str(results_ws["J4"].value)
        evidence_detail = str(evidence_ws["D4"].value)
        assert '"dhcp_lease_acknowledged": true' in evidence_detail.lower()
        assert '"offered_ip": "192.168.4.68"' in evidence_detail
    finally:
        wb.close()


@pytest.mark.asyncio
async def test_excel_export_strips_template_trash_entries(tmp_path, monkeypatch):
    import zipfile

    monkeypatch.setattr(report_generator.settings, "REPORT_DIR", str(tmp_path))

    output = await report_generator.generate_excel_report(
        _make_test_run(),
        [_make_test_result()],
        template_key="generic",
    )

    with zipfile.ZipFile(output) as zf:
        names = [i.filename for i in zf.infolist()]

    assert not any(n.startswith("[trash]/") for n in names), \
        f"Stray [trash]/ entries survived patching: {names}"


@pytest.mark.asyncio
async def test_engineer_notes_written_to_excel_cell(tmp_path, monkeypatch):
    monkeypatch.setattr(report_generator.settings, "REPORT_DIR", str(tmp_path))
    note = "Manual reboot observation"

    output = await report_generator.generate_excel_report(
        _make_test_run(),
        [_make_test_result_with_notes(note)],
        template_key="generic",
    )

    wb = load_workbook(output)
    try:
        ws = wb["Test Results"]
        header = ws["I3"].value
        assert header == "Engineer Notes"
        column_values = [ws[f"I{row}"].value for row in range(4, 10)]
        assert note in column_values
    finally:
        wb.close()


def test_build_report_document_includes_readiness_summary():
    report = report_generator.build_report_document(
        _make_test_run(),
        [_make_test_result()],
        template_key="generic",
        readiness_summary={
            "score": 9,
            "level": "conditional",
            "label": "Operational with advisories",
            "report_ready": True,
            "operational_ready": False,
            "blocking_issue_count": 0,
            "pending_manual_count": 0,
            "release_blocking_failure_count": 0,
            "review_required_issue_count": 0,
            "manual_evidence_pending_count": 0,
            "advisory_count": 1,
            "override_count": 0,
            "failed_test_count": 0,
            "completed_result_count": 1,
            "total_result_count": 1,
            "trust_tier_counts": {
                "release_blocking": 1,
                "review_required": 0,
                "advisory": 0,
                "manual_evidence": 0,
            },
            "next_step": "Issue the report with the advisory notes and follow-up actions captured.",
            "reasons": ["1 advisory finding should be called out in the report."],
            "summary": "Operational with advisories (9/10). 1 advisory finding should be called out in the report.",
        },
    )

    assert report.readiness_summary["score"] == 9
    assert "Readiness status: Operational with advisories (9/10)." in report.metadata["summary_text"]
    assert "Official Report Ready: Yes" in report.additional_sections[1].body
