"""Report generation engine."""

from __future__ import annotations

import csv
import json
import logging
import os
import re
import shutil
import subprocess
import textwrap
import uuid
from dataclasses import dataclass, field
from datetime import datetime, timezone
from functools import lru_cache
from pathlib import Path
from typing import Any, Optional

from app.config import settings
from app.services.run_readiness import build_run_readiness_summary
from app.services.scenario_routing import describe_connection_scenario
from app.utils.datetime import as_utc, utcnow_naive

logger = logging.getLogger(__name__)

_ILLEGAL_XML_CHARS = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]")
_MAPPINGS_DIR = Path(__file__).resolve().parent / "cell_mappings"
_ELECTRACOM_REPORT_LOGO = "electracom-logo.png"

TEMPLATE_FILES = {
    "pelco_camera": "1TS - Pelco SMLE1-15V5-3H Camera Device Qualification Rev 2.xlsx",
    "easyio_controller": "EasyIO FW08 - Device Testing Plan - v1.1.xlsx",
    "sauter_680_as": "Sauter - 680-AS - IP Device Qualification Template C00.xlsx",
    "generic": "[MANUFACTURER] - [MODEL] - IP Device Qualification Template C00 - ADDED TESTING Scenarios.xlsx",
}
TEMPLATE_MAPPING_FILES = {
    "pelco_camera": "pelco_camera.json",
    "easyio_controller": "easyio_controller.json",
    "sauter_680_as": "sauter_680_as.json",
    "generic": "generic.json",
}
ACTIVE_REPORT_TEMPLATE_KEYS = ("generic",)
REPORT_EXTENSIONS = {".xlsx", ".docx", ".pdf", ".csv", ".dxf"}
REPORT_OUTPUT_FILENAME_RE = re.compile(
    r"^EDQ_Report_"
    r"[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}_"
    r"(?:generic|pelco_camera|easyio_controller|sauter_680_as)_"
    r"\d{8}_\d{6}\.(?:xlsx|docx)$",
    re.IGNORECASE,
)

TEMPLATE_INFO = {
    "pelco_camera": {"name": "Pelco Camera (Rev 2)", "device_category": "camera", "description": "Pelco camera workbook."},
    "easyio_controller": {"name": "EasyIO Controller", "device_category": "controller", "description": "EasyIO controller workbook."},
    "sauter_680_as": {"name": "Sauter 680-AS (C00)", "device_category": "controller", "description": "Sauter 680-AS qualification workbook."},
    "generic": {"name": "Generic IP Device (Rev00 C00)", "device_category": "generic", "description": "Canonical four-section qualification report."},
}


def _resolve_templates_dir() -> Path:
    here = Path(__file__).resolve()
    candidates: list[Path] = []

    env_value = os.getenv("EDQ_TEMPLATES_DIR")
    if env_value:
        candidates.append(Path(env_value))

    for depth in (2, 4):
        try:
            candidates.append(here.parents[depth] / "templates")
        except IndexError:
            continue

    candidates.append(Path.cwd() / "templates")

    for candidate in candidates:
        if candidate.exists():
            return candidate

    return candidates[0]


_TEMPLATES_DIR = _resolve_templates_dir()


@lru_cache(maxsize=1)
def _repo_root() -> Path:
    here = Path(__file__).resolve()
    for parent in here.parents:
        if (parent / "AGENTS.md").exists() or (parent / "docker-compose.yml").exists():
            return parent
    return here.parents[2]


def _resolve_existing_path(candidates: list[Path]) -> Path | None:
    for candidate in candidates:
        try:
            resolved = candidate.expanduser().resolve()
        except OSError:
            continue
        if resolved.is_file():
            return resolved
    return None


@lru_cache(maxsize=1)
def _resolve_electracom_report_logo_path() -> Path | None:
    root = _repo_root()
    backend_root = Path(__file__).resolve().parents[2]
    return _resolve_existing_path(
        [
            root / "frontend" / "public" / _ELECTRACOM_REPORT_LOGO,
            Path.cwd() / "assets" / _ELECTRACOM_REPORT_LOGO,
            root / "assets" / _ELECTRACOM_REPORT_LOGO,
            backend_root / "assets" / _ELECTRACOM_REPORT_LOGO,
        ]
    )


def _resolve_uploaded_logo_path(logo_path: str) -> Path | None:
    if not logo_path:
        return None
    raw = Path(logo_path)
    if raw.is_absolute():
        candidates = [raw]
    else:
        filename = raw.name
        candidates = [
            Path(settings.UPLOAD_DIR) / "branding" / filename,
            Path(settings.UPLOAD_DIR) / filename,
        ]
        if str(raw) == filename:
            candidates.append(Path.cwd() / filename)
    return _resolve_existing_path(candidates)


def _scaled_logo_pixels(logo_path: Path, width_px: int = 190) -> tuple[int, int]:
    try:
        from PIL import Image

        with Image.open(logo_path) as image:
            source_width, source_height = image.size
        if source_width <= 0 or source_height <= 0:
            return width_px, 35
        return width_px, max(1, round(width_px * source_height / source_width))
    except Exception:
        return width_px, 35

_VERDICT_MAP = {
    "pass": "PASS",
    "qualified_pass": "QUALIFIED PASS",
    "fail": "FAIL",
    "advisory": "ADVISORY",
    "na": "N/A",
    "info": "INFO",
    "error": "ERROR",
    "pending": "PENDING",
    "running": "RUNNING",
    "incomplete": "INCOMPLETE",
}
_SUMMARY_FIELDS = [
    ("test_attempt", "Test Attempt"),
    ("date_range", "Date Test Started - Date Test Finished"),
    ("system", "System"),
    ("system_owner", "System Owner"),
    ("manufacturer", "Manufacturer"),
    ("model", "Model"),
    ("firmware", "Firmware Version"),
    ("serial", "Serial Number"),
    ("tester_name", "Name Of Tester"),
    ("overall_result", "TEST RESULT"),
]
_SUMMARY_LABELS = {
    **dict(_SUMMARY_FIELDS),
    "start_date": "Date Test Started",
    "end_date": "Date Test Finished",
}
_TESTPLAN_COLUMN_LABELS = {
    "test_number": "Test Number",
    "brief_description": "Brief Description",
    "test_description": "Test Description",
    "tier": "Tier",
    "tool": "Tool",
    "essential_test": "Essential Test",
    "essential_pass": "Essential Pass",
    "test_result": "Test Result",
    "test_comments": "Test Comments",
    "engineer_notes": "Engineer Notes",
    "evidence_summary": "Evidence Summary",
    "raw_evidence": "Raw Evidence",
    "script_flag": "Script Flag",
}


@dataclass
class ReportRow:
    test_number: str
    brief_description: str
    test_description: str
    essential_test: str
    test_result: str
    test_comments: str
    tier: str = ""
    tool: str = ""
    engineer_notes: str = ""
    evidence_summary: str = ""
    evidence_detail: str = ""
    script_flag: str = ""
    template_backed: bool = False


@dataclass
class ReportSection:
    title: str
    body: str


@dataclass
class ReportBranding:
    company_name: str = ""
    primary_color: str = ""
    footer_text: str = ""
    logo_path: str = ""


@dataclass
class ReportDocument:
    template_key: str
    template_path: Path
    mapping: dict[str, Any]
    metadata: dict[str, str]
    summary_section_title: str = "TEST SUMMARY"
    testplan_section_title: str = "TESTPLAN"
    additional_section_title: str = "ADDITIONAL INFORMATION"
    summary_text_label: str = "Summary"
    summary_fields: list[tuple[str, str]] = field(default_factory=list)
    testplan_columns: list[tuple[str, str]] = field(default_factory=list)
    branding: ReportBranding = field(default_factory=ReportBranding)
    readiness_summary: dict[str, Any] = field(default_factory=dict)
    rows: list[ReportRow] = field(default_factory=list)
    additional_sections: list[ReportSection] = field(default_factory=list)


def _sanitize(value: Any) -> Any:
    return _ILLEGAL_XML_CHARS.sub("", value) if isinstance(value, str) else value


def _sanitize_text(value: Any) -> str:
    if value is None:
        return ""
    if hasattr(value, "value"):
        value = value.value
    return str(_sanitize(str(value)))


def _safe_attr(obj: Any, attr: str, default: str = "") -> str:
    value = getattr(obj, attr, None) if obj is not None else None
    if hasattr(value, "value"):
        value = value.value
    return _sanitize_text(default if value is None else value)


def _validate_template_key(template_key: str) -> str:
    if template_key not in TEMPLATE_FILES:
        raise ValueError("Unknown template key")
    return template_key


def _validate_report_extension(extension: str) -> str:
    if extension not in REPORT_EXTENSIONS:
        raise ValueError("Unsupported report extension")
    return extension


def _safe_test_run_id(test_run: Any) -> str:
    return str(uuid.UUID(str(test_run.id)))


def _report_source_file(source_path: str) -> Path:
    source_name = str(source_path).replace("\\", "/").rsplit("/", 1)[-1]
    if not REPORT_OUTPUT_FILENAME_RE.fullmatch(source_name):
        raise RuntimeError("Invalid report source filename")

    output_dir = Path(settings.REPORT_DIR).resolve()
    for candidate in output_dir.iterdir():
        if candidate.name == source_name and candidate.is_file():
            resolved = candidate.resolve()
            if resolved.is_relative_to(output_dir):
                return resolved
    raise RuntimeError("Report source file not found")


def _load_mapping(template_key: str) -> dict[str, Any]:
    safe_template_key = _validate_template_key(template_key)
    path = _MAPPINGS_DIR / TEMPLATE_MAPPING_FILES[safe_template_key]
    return json.loads(path.read_text(encoding="utf-8")) if path.exists() else {}


def _json_dict(value: Any) -> dict[str, Any]:
    if isinstance(value, dict):
        return value
    if isinstance(value, str):
        try:
            parsed = json.loads(value)
        except (json.JSONDecodeError, TypeError):
            return {}
        return parsed if isinstance(parsed, dict) else {}
    return {}


def _resolve_branding(report_config: Any = None, branding_settings: Any = None) -> ReportBranding:
    config_branding = _json_dict(getattr(report_config, "branding", None))
    company_name = (
        _safe_attr(branding_settings, "company_name")
        or _safe_attr(report_config, "client_name")
        or str(config_branding.get("company_name") or "")
        or settings.APP_NAME
    )
    primary_color = (
        _safe_attr(branding_settings, "primary_color")
        or str(config_branding.get("primary_color") or "")
        or "#2563eb"
    )
    footer_text = (
        _safe_attr(branding_settings, "footer_text")
        or str(config_branding.get("footer_text") or "")
    )
    logo_path = (
        _safe_attr(branding_settings, "logo_path")
        or _safe_attr(report_config, "logo_path")
        or str(config_branding.get("logo_path") or "")
    )
    return ReportBranding(
        company_name=_sanitize_text(company_name),
        primary_color=_sanitize_text(primary_color),
        footer_text=_sanitize_text(footer_text),
        logo_path=_sanitize_text(logo_path),
    )


def _report_window(test_run: Any) -> tuple[Any, Any]:
    start = getattr(test_run, "started_at", None) or getattr(test_run, "created_at", None)
    end = getattr(test_run, "completed_at", None) or utcnow_naive()
    return start, end


def _format_date_range(test_run: Any) -> str:
    start, end = _report_window(test_run)
    return f"{start.strftime('%d/%m/%Y') if start else ''} - {end.strftime('%d/%m/%Y') if end else ''}".strip(" -")


def _overall_result(test_run: Any) -> str:
    raw = getattr(test_run, "overall_verdict", None)
    raw = raw.value if hasattr(raw, "value") else str(raw or "incomplete")
    return _VERDICT_MAP.get(raw, raw.upper())


def _comment(result: Any) -> str:
    parts = [getattr(result, "comment_override", None) or getattr(result, "comment", None) or ""]
    reason = getattr(result, "override_reason", None)
    if reason:
        parts.append(f"[Override: {reason}]")
    return _sanitize_text(" ".join(str(part) for part in parts if part)).strip()


def _engineer_notes(result: Any) -> str:
    notes = getattr(result, "engineer_notes", None)
    return _sanitize_text(notes) if notes else ""


def _clip_text(value: Any, limit: int = 2000) -> str:
    text = _sanitize_text(value).strip()
    if not text:
        return ""
    if len(text) <= limit:
        return text
    return f"{text[:limit].rstrip()} ... [truncated]"


def _format_evidence(result: Any, limit: int | None = None) -> str:
    """Format evidence from a result object with optional truncation.
    
    Checks raw_output → findings → parsed_data in order.
    For findings and parsed_data, attempts JSON serialization with proper formatting.
    Falls back to string representation on TypeError.
    Applies truncation only when limit is provided.
    """
    raw_output = getattr(result, "raw_output", None)
    if raw_output:
        text = _sanitize_text(raw_output).strip()
        return _clip_text(text, limit=limit) if limit else text

    findings = getattr(result, "findings", None)
    if findings:
        try:
            text = json.dumps(findings, indent=2, ensure_ascii=False)
        except TypeError:
            text = str(findings).strip()
        text = _sanitize_text(text).strip()
        return _clip_text(text, limit=limit) if limit else text

    parsed = getattr(result, "parsed_data", None)
    if parsed:
        try:
            text = json.dumps(parsed, indent=2, ensure_ascii=False)
        except TypeError:
            text = str(parsed).strip()
        text = _sanitize_text(text).strip()
        return _clip_text(text, limit=limit) if limit else text

    return ""


def _format_tier(result: Any) -> str:
    raw = _safe_attr(result, "tier").replace("_", " ").strip()
    return raw.title() if raw else ""


def _format_tool(result: Any) -> str:
    return _safe_attr(result, "tool")


def _evidence_summary(result: Any) -> str:
    return _format_evidence(result, limit=2500)


def _evidence_detail(result: Any) -> str:
    return _format_evidence(result, limit=None)


_VERDICT_RANK = {
    "": 0,
    "na": 10,
    "pass": 20,
    "qualified_pass": 30,
    "info": 40,
    "advisory": 50,
    "running": 60,
    "pending": 60,
    "incomplete": 70,
    "error": 80,
    "fail": 90,
}


def _aggregate_verdict(results: list[Any]) -> str:
    if not results:
        return ""
    selected = max(
        results,
        key=lambda result: _VERDICT_RANK.get(_safe_attr(result, "verdict").lower(), 0),
    )
    verdict = _safe_attr(selected, "verdict")
    return _VERDICT_MAP.get(verdict.lower(), verdict.upper())


def _combine_source_text(results: list[Any], formatter: Any) -> str:
    parts: list[str] = []
    for result in results:
        text = _sanitize_text(formatter(result)).strip()
        if not text:
            continue
        if len(results) > 1:
            test_id = _safe_attr(result, "test_id") or "source"
            parts.append(f"{test_id}: {text}")
        else:
            parts.append(text)
    return "\n".join(parts)


def _template_comment(
    results: list[Any],
    has_engineer_notes_column: bool,
    has_evidence_column: bool,
) -> str:
    if not results:
        return ""
    parts = [_combine_source_text(results, _comment)]
    if not has_engineer_notes_column:
        notes = _combine_source_text(results, _engineer_notes)
        if notes:
            parts.append(f"Engineer Notes:\n{notes}")
    if not has_evidence_column:
        evidence = _combine_source_text(results, _evidence_summary)
        if evidence:
            parts.append(f"Evidence:\n{evidence}")
    return _sanitize_text("\n".join(part for part in parts if part).strip())


def _test_sort_key(result: Any) -> tuple[int, str]:
    test_id = _safe_attr(result, "test_id")
    match = re.match(r"^[A-Za-z]+(\d+)$", test_id)
    if match:
        return (int(match.group(1)), test_id)
    return (9999, test_id)


def _generic_testplan_columns() -> list[tuple[str, str]]:
    return [
        ("test_number", "Test Number"),
        ("brief_description", "Brief Description"),
        ("test_description", "Test Description"),
        ("essential_test", "Essential Test"),
        ("test_result", "Test Result"),
        ("test_comments", "Test Comments"),
        ("tier", "Tier"),
        ("tool", "Tool"),
        ("engineer_notes", "Engineer Notes"),
        ("evidence_summary", "Evidence Summary"),
        ("evidence_detail", "Raw Evidence"),
    ]


def _generic_template_report_columns(mapping: dict[str, Any], template_path: Path) -> list[tuple[str, str]]:
    if not template_path.exists() or not mapping.get("row_sources"):
        return _generic_testplan_columns()

    columns = _resolve_testplan_columns(mapping)
    attributes = {attribute for attribute, _ in columns}
    if "engineer_notes" not in attributes:
        columns.append(("engineer_notes", "Engineer Notes"))
    if "evidence_summary" not in attributes:
        columns.append(("evidence_summary", "Evidence Summary"))
    return columns


def _build_detailed_report_rows(test_results: list[Any]) -> list[ReportRow]:
    rows: list[ReportRow] = []
    for result in sorted(test_results, key=_test_sort_key):
        verdict = _safe_attr(result, "verdict")
        rows.append(
            ReportRow(
                test_number=_safe_attr(result, "test_id"),
                brief_description=_safe_attr(result, "test_name"),
                test_description="",
                essential_test=_safe_attr(result, "is_essential").upper(),
                test_result=_VERDICT_MAP.get(verdict.lower(), verdict.upper()),
                test_comments=_comment(result),
                tier=_format_tier(result),
                tool=_format_tool(result),
                engineer_notes=_engineer_notes(result),
                evidence_summary=_evidence_summary(result),
                evidence_detail=_evidence_detail(result),
            )
        )
    return rows


def _summary_text(
    test_run: Any,
    metadata: dict[str, str],
    include_synopsis: bool,
    readiness_summary: dict[str, Any],
) -> str:
    synopsis = _sanitize_text(getattr(test_run, "synopsis", None)).replace("[AI-DRAFTED] ", "").strip()
    readiness_line = (
        f" Readiness status: {readiness_summary.get('label', 'Unknown')} "
        f"({readiness_summary.get('score', 1)}/10)."
    )
    if include_synopsis and synopsis:
        return _sanitize_text(
            f"{synopsis.rstrip('.')}.{readiness_line}" if not synopsis.endswith(".") else f"{synopsis}{readiness_line}"
        )
    summary = (
        f"Qualification testing for {metadata.get('manufacturer') or 'Unknown manufacturer'} "
        f"{metadata.get('model') or 'Unknown model'} completed with an overall result of "
        f"{metadata.get('overall_result') or 'INCOMPLETE'}."
    )
    return _sanitize_text(f"{summary}{readiness_line}")


def _supporting_evidence_body(
    test_run: Any,
    branding: ReportBranding,
    readiness_summary: dict[str, Any],
) -> str:
    parts = [
        f"Readiness Status: {readiness_summary.get('label', 'Unknown')} ({readiness_summary.get('score', 1)}/10)",
        f"Official Report Ready: {'Yes' if readiness_summary.get('report_ready') else 'No'}",
        f"Operationally Ready: {'Yes' if readiness_summary.get('operational_ready') else 'No'}",
        f"Next Step: {readiness_summary.get('next_step', 'Review the remaining findings.')}",
        "Key Reasons:",
    ]
    for reason in readiness_summary.get("reasons", [])[:4]:
        parts.append(f"- {reason}")
    parts.append(
        f"Connection Scenario: {describe_connection_scenario(_safe_attr(test_run, 'connection_scenario', 'direct'))}"
    )
    if branding.footer_text:
        parts.append(f"Report Footer: {branding.footer_text}")
    return _sanitize_text("\n".join(parts))


def _resolve_summary_fields(mapping: dict[str, Any]) -> list[tuple[str, str]]:
    metadata_cells = mapping.get("metadata_cells") or {}
    ordered_keys = [key for key in metadata_cells if key not in {"summary_text", "synopsis_text"}]
    if not ordered_keys:
        return list(_SUMMARY_FIELDS)
    return [(key, _SUMMARY_LABELS.get(key, key.replace("_", " ").title())) for key in ordered_keys]


def _resolve_testplan_columns(mapping: dict[str, Any]) -> list[tuple[str, str]]:
    columns = mapping.get("testplan_columns") or {}
    ordered_keys = [
        "test_number",
        "brief_description",
        "test_description",
        "essential_test",
        "essential_pass",
        "test_result",
        "test_comments",
        "tier",
        "tool",
        "engineer_notes",
        "evidence_summary",
        "raw_evidence",
        "script_flag",
    ]
    resolved: list[tuple[str, str]] = []
    for key in ordered_keys:
        if key not in columns:
            continue
        if key == "essential_pass":
            attribute = "essential_test"
        elif key == "raw_evidence":
            attribute = "evidence_detail"
        else:
            attribute = key
        resolved.append((attribute, _TESTPLAN_COLUMN_LABELS.get(key, key.replace("_", " ").title())))
    if resolved:
        return resolved
    return [
        ("test_number", "Test Number"),
        ("brief_description", "Brief Description"),
        ("test_description", "Test Description"),
        ("essential_test", "Essential Test"),
        ("test_result", "Test Result"),
        ("test_comments", "Test Comments"),
        ("engineer_notes", "Engineer Notes"),
    ]


def _resolve_summary_text_label(mapping: dict[str, Any]) -> str:
    metadata_cells = mapping.get("metadata_cells") or {}
    return "Synopsis" if "synopsis_text" in metadata_cells else "Summary"


def _report_row_values(row: ReportRow, columns: list[tuple[str, str]]) -> list[str]:
    return [_sanitize_text(getattr(row, attribute, "")) for attribute, _ in columns]


def _pdf_safe_width(pdf: Any) -> float:
    return max(float(pdf.w) - float(pdf.l_margin) - float(pdf.r_margin), 20.0)


def _pdf_safe_text(value: str) -> str:
    text = _sanitize_text(value)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("|", " - ")
    text = (
        text.replace("\u2018", "'")
        .replace("\u2019", "'")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
        .replace("\u2013", "-")
        .replace("\u2014", "-")
        .replace("\u2022", "-")
        .replace("\u2026", "...")
        .replace("\u00a0", " ")
    )
    return text.encode("latin-1", "replace").decode("latin-1")


def _pdf_wrapped_lines(value: str, width: int = 140) -> list[str]:
    lines: list[str] = []
    for paragraph in _pdf_safe_text(value).split("\n"):
        wrapped = textwrap.wrap(
            paragraph,
            width=width,
            break_long_words=True,
            break_on_hyphens=True,
        )
        lines.extend(wrapped or [""])
    return lines or [""]


def _write_pdf_lines(pdf: Any, value: str, line_height: float = 5, wrap_width: int = 140) -> None:
    width = _pdf_safe_width(pdf)
    for line in _pdf_wrapped_lines(value, width=wrap_width):
        pdf.set_x(pdf.l_margin)
        pdf.cell(width, line_height, _pdf_safe_text(line or " "), new_x="LMARGIN", new_y="NEXT")


def build_report_document(
    test_run: Any,
    test_results: list[Any],
    report_config: Any = None,
    template_key: str = "generic",
    enabled_test_ids: Optional[list[str]] = None,
    whitelist_entries: Optional[list[Any]] = None,
    include_synopsis: bool = False,
    branding_settings: Any = None,
    readiness_summary: Optional[dict[str, Any]] = None,
) -> ReportDocument:
    del whitelist_entries
    template_key = _validate_template_key(template_key)
    enabled_ids = set(enabled_test_ids or [])
    filtered_results = sorted(
        [r for r in test_results if not enabled_ids or getattr(r, "test_id", None) in enabled_ids],
        key=_test_sort_key,
    )
    mapping = _load_mapping(template_key)
    template_path = _TEMPLATES_DIR / TEMPLATE_FILES[template_key]
    branding = _resolve_branding(report_config, branding_settings)
    start, end = _report_window(test_run)
    resolved_readiness = readiness_summary or build_run_readiness_summary(test_run, filtered_results)
    summary_section_title = str(mapping.get("synopsis_sheet") or "TEST SUMMARY")
    testplan_section_title = str(mapping.get("testplan_sheet") or "TESTPLAN")
    additional_section_title = str(mapping.get("additional_sheet") or "ADDITIONAL INFORMATION")
    summary_text_label = _resolve_summary_text_label(mapping)
    summary_fields = _resolve_summary_fields(mapping)
    testplan_columns = (
        _generic_template_report_columns(mapping, template_path)
        if template_key == "generic"
        else _resolve_testplan_columns(mapping)
    )

    device = getattr(test_run, "device", None)
    engineer = getattr(test_run, "engineer", None)
    metadata = {
        "test_attempt": "1",
        "date_range": _format_date_range(test_run),
        "start_date": start.strftime("%d/%m/%Y") if start else "",
        "end_date": end.strftime("%d/%m/%Y") if end else "",
        "system": _safe_attr(device, "category").replace("_", " ").title(),
        "system_owner": _safe_attr(report_config, "client_name") or branding.company_name,
        "manufacturer": _safe_attr(device, "manufacturer"),
        "model": _safe_attr(device, "model"),
        "firmware": _safe_attr(device, "firmware_version"),
        "serial": _safe_attr(device, "serial_number"),
        "tester_name": _safe_attr(engineer, "full_name"),
        "overall_result": _overall_result(test_run),
        "summary_text": "",
    }
    metadata["summary_text"] = _summary_text(
        test_run,
        metadata,
        include_synopsis,
        resolved_readiness,
    )
    metadata["synopsis_text"] = metadata["summary_text"]

    rows: list[ReportRow] = []
    if template_path.exists() and mapping and mapping.get("row_sources"):
        from openpyxl import load_workbook

        wb = load_workbook(str(template_path), read_only=True, data_only=False)
        try:
            ws = wb[mapping["testplan_sheet"]]
            cols = mapping["testplan_columns"]
            start = mapping["testplan_start_row"]
            count = mapping["testplan_row_count"]
            result_by_id = {str(getattr(r, "test_id", "")): r for r in filtered_results}
            sources = mapping.get("row_sources", {})
            brief_column = cols.get("brief_description") or cols.get("test_description")
            description_column = cols.get("test_description") or brief_column
            essential_column = cols.get("essential_test", cols.get("essential_pass"))
            engineer_notes_column = cols.get("engineer_notes")
            evidence_column = cols.get("evidence_summary") or cols.get("evidence_detail")
            script_column = cols.get("script_flag")
            for row_index in range(start, start + count):
                number = str(ws[f"{cols['test_number']}{row_index}"].value or "").strip()
                source_ids = sources.get(number, [])
                source_results = [result_by_id[test_id] for test_id in source_ids if test_id in result_by_id]
                rows.append(
                    ReportRow(
                        test_number=number,
                        brief_description=str(ws[f"{brief_column}{row_index}"].value or "") if brief_column else "",
                        test_description=str(ws[f"{description_column}{row_index}"].value or "") if description_column else "",
                        essential_test=str(ws[f"{essential_column}{row_index}"].value or "") if essential_column else "",
                        test_result=_aggregate_verdict(source_results),
                        test_comments=_template_comment(
                            source_results,
                            has_engineer_notes_column=bool(engineer_notes_column),
                            has_evidence_column=bool(evidence_column),
                        ),
                        engineer_notes=_combine_source_text(source_results, _engineer_notes),
                        evidence_summary=_combine_source_text(source_results, _evidence_summary),
                        evidence_detail=_combine_source_text(source_results, _evidence_detail),
                        script_flag=str(ws[f"{script_column}{row_index}"].value or "") if script_column else "",
                        template_backed=True,
                    )
                )
        finally:
            wb.close()
    elif template_key == "generic":
        rows = _build_detailed_report_rows(filtered_results)
    else:
        rows = _build_detailed_report_rows(filtered_results)

    additional_sections = [
        ReportSection("Executive Summary", metadata["summary_text"]),
        ReportSection(
            "Readiness and Supporting Evidence",
            _supporting_evidence_body(test_run, branding, resolved_readiness),
        ),
    ]
    return ReportDocument(
        template_key=template_key,
        template_path=template_path,
        mapping=mapping,
        metadata=metadata,
        summary_section_title=summary_section_title,
        testplan_section_title=testplan_section_title,
        additional_section_title=additional_section_title,
        summary_text_label=summary_text_label,
        summary_fields=summary_fields,
        testplan_columns=testplan_columns,
        branding=branding,
        readiness_summary=resolved_readiness,
        rows=rows,
        additional_sections=additional_sections,
    )


def get_available_templates() -> list[dict[str, Any]]:
    items = []
    for key in ACTIVE_REPORT_TEMPLATE_KEYS:
        info = TEMPLATE_INFO[key]
        items.append({
            "key": key,
            "name": info["name"],
            "device_category": info["device_category"],
            "description": info["description"],
            "template_exists": True if key == "generic" else (_TEMPLATES_DIR / TEMPLATE_FILES[key]).exists(),
            "mapping_exists": (_MAPPINGS_DIR / TEMPLATE_MAPPING_FILES[key]).exists(),
        })
    return items


def _output_path(test_run: Any, template_key: str, extension: str) -> Path:
    output_dir = Path(settings.REPORT_DIR).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    run_id = _safe_test_run_id(test_run)
    safe_template_key = _validate_template_key(template_key)
    safe_extension = _validate_report_extension(extension)
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    return output_dir / f"EDQ_Report_{run_id}_{safe_template_key}_{stamp}{safe_extension}"


def _add_openpyxl_logo(ws: Any, logo_path: Path | None, cell: str, width_px: int = 190) -> None:
    if logo_path is None:
        return
    try:
        from openpyxl.drawing.image import Image as ExcelImage

        width, height = _scaled_logo_pixels(logo_path, width_px)
        image = ExcelImage(str(logo_path))
        image.width = width
        image.height = height
        ws.add_image(image, cell)
        match = re.match(r"^[A-Z]+(\d+)$", cell, re.IGNORECASE)
        if match:
            row_number = int(match.group(1))
            ws.row_dimensions[row_number].height = max(
                ws.row_dimensions[row_number].height or 0,
                height * 0.75 + 4,
            )
    except Exception as exc:
        logger.warning("Failed to add report logo to worksheet %s: %s", ws.title, exc)


def _logo_display_name(path: Path | None) -> str:
    return path.name if path else ""


def _write_generic_excel_report(path: Path, report: ReportDocument) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    summary_ws = wb.active
    summary_ws.title = report.summary_section_title
    results_ws = wb.create_sheet(report.testplan_section_title)
    additional_ws = wb.create_sheet(report.additional_section_title)

    title_fill = PatternFill("solid", fgColor="1F4E78")
    band_fill = PatternFill("solid", fgColor="D9EAF7")
    title_font = Font(color="FFFFFF", bold=True, size=14)
    header_font = Font(bold=True)
    mono_font = Font(name="Consolas")
    wrap = Alignment(vertical="top", wrap_text=True)
    issuer_logo = _resolve_electracom_report_logo_path()
    client_logo = _resolve_uploaded_logo_path(report.branding.logo_path)

    _add_openpyxl_logo(summary_ws, issuer_logo, "A1", width_px=190)
    if client_logo:
        _add_openpyxl_logo(summary_ws, client_logo, "G1", width_px=150)
    summary_ws.merge_cells("C1:F1" if client_logo else "C1:H1")
    summary_ws["C1"] = "IP Device Qualification Report"
    summary_ws["C1"].fill = title_fill
    summary_ws["C1"].font = title_font
    summary_ws["C1"].alignment = Alignment(horizontal="center", vertical="center")
    summary_ws["A2"] = f"Template Profile: {TEMPLATE_INFO[report.template_key]['name']}"

    row = 4
    for key, label in report.summary_fields:
        summary_ws[f"A{row}"] = label
        summary_ws[f"A{row}"].font = header_font
        summary_ws[f"A{row}"].fill = band_fill
        summary_ws[f"B{row}"] = report.metadata.get(key, "")
        summary_ws[f"B{row}"].alignment = wrap
        row += 1

    summary_ws[f"A{row + 1}"] = report.summary_text_label
    summary_ws[f"A{row + 1}"].font = header_font
    summary_ws[f"A{row + 1}"].fill = band_fill
    summary_ws[f"B{row + 1}"] = report.metadata.get("summary_text", "")
    summary_ws[f"B{row + 1}"].alignment = wrap
    summary_ws.column_dimensions["A"].width = 28
    summary_ws.column_dimensions["B"].width = 110
    summary_ws.freeze_panes = "A4"

    _add_openpyxl_logo(results_ws, issuer_logo, "A1", width_px=150)
    start_col_idx = 2
    end_col_idx = start_col_idx + len(report.testplan_columns) - 1
    start_col = get_column_letter(start_col_idx)
    end_col = get_column_letter(end_col_idx)
    results_ws.merge_cells(f"{start_col}1:{end_col}1")
    results_ws["B1"] = "Per-Test Results and Evidence"
    results_ws["B1"].fill = title_fill
    results_ws["B1"].font = title_font
    results_ws["B1"].alignment = Alignment(horizontal="center", vertical="center")
    header_row = 3
    columns = [get_column_letter(idx) for idx in range(start_col_idx, end_col_idx + 1)]
    width_hints = {
        "test_number": (12, 14),
        "brief_description": (22, 34),
        "test_description": (36, 64),
        "essential_test": (14, 16),
        "test_result": (14, 18),
        "test_comments": (44, 70),
        "tier": (14, 20),
        "tool": (12, 18),
        "engineer_notes": (24, 40),
        "evidence_summary": (42, 72),
        "evidence_detail": (55, 95),
    }

    def _suggest_width(attribute: str, label: str, values: list[str]) -> float:
        min_width, max_width = width_hints.get(attribute, (12, 40))
        longest = max([len(label), *(min(len(line), max_width) for value in values for line in str(value).splitlines())], default=min_width)
        return float(max(min_width, min(max_width, longest + 2)))

    row_values = [
        _report_row_values(row_values, report.testplan_columns)
        for row_values in report.rows
    ]
    values_by_attribute = {
        attribute: [values[idx] for values in row_values if idx < len(values)]
        for idx, (attribute, _label) in enumerate(report.testplan_columns)
    }

    for col, (attribute, label) in zip(columns, report.testplan_columns, strict=False):
        results_ws[f"{col}{header_row}"] = label
        results_ws[f"{col}{header_row}"].font = header_font
        results_ws[f"{col}{header_row}"].fill = band_fill
        results_ws[f"{col}{header_row}"].alignment = wrap
        results_ws.column_dimensions[col].width = _suggest_width(attribute, label, values_by_attribute.get(attribute, []))

    for offset, values in enumerate(row_values, start=header_row + 1):
        for col, (attribute, _label), value in zip(columns, report.testplan_columns, values, strict=False):
            results_ws[f"{col}{offset}"] = value
            results_ws[f"{col}{offset}"].alignment = wrap
            if attribute == "evidence_detail" or "\nPORT" in str(value) or str(value).startswith("PORT"):
                results_ws[f"{col}{offset}"].font = mono_font

    results_ws.freeze_panes = "B4"
    results_ws.auto_filter.ref = f"B{header_row}:{end_col}{max(header_row, len(report.rows) + header_row)}"

    additional_ws.merge_cells("A1:B1")
    additional_ws["A1"] = "Supporting Notes"
    additional_ws["A1"].fill = title_fill
    additional_ws["A1"].font = title_font
    additional_ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    additional_ws.column_dimensions["A"].width = 28
    additional_ws.column_dimensions["B"].width = 110
    section_row = 3
    for section in report.additional_sections:
        additional_ws[f"A{section_row}"] = section.title
        additional_ws[f"A{section_row}"].font = header_font
        additional_ws[f"A{section_row}"].fill = band_fill
        additional_ws[f"B{section_row}"] = section.body
        additional_ws[f"B{section_row}"].alignment = wrap
        section_row += 2

    wb.save(str(path))


async def generate_excel_report(
    test_run: Any,
    test_results: list[Any],
    report_config: Any = None,
    template_key: str = "generic",
    enabled_test_ids: Optional[list[str]] = None,
    whitelist_entries: Optional[list[Any]] = None,
    include_synopsis: bool = False,
    branding_settings: Any = None,
    readiness_summary: Optional[dict[str, Any]] = None,
) -> str:
    import asyncio

    report = build_report_document(
        test_run,
        test_results,
        report_config,
        template_key,
        enabled_test_ids,
        whitelist_entries,
        include_synopsis,
        branding_settings,
        readiness_summary,
    )
    path = _output_path(test_run, template_key, ".xlsx")

    if template_key == "generic":
        await asyncio.to_thread(_write_generic_excel_report, path, report)
    elif report.template_path.exists() and report.mapping:
        # Build cell updates per sheet — ZIP-level patcher preserves ALL
        # template assets (images, drawings, printer settings, styles, etc.)
        from app.services.xlsx_template_patcher import XlsxImageInsert, patch_xlsx

        sheet_updates: dict[str, dict[str, str | None]] = {}

        # TEST SUMMARY metadata cells
        synopsis_sheet = report.mapping["synopsis_sheet"]
        summary_cells: dict[str, str | None] = {}
        for key, cell in report.mapping.get("metadata_cells", {}).items():
            value = _sanitize(report.metadata.get(key, ""))
            summary_cells[cell] = value if value else None
        if summary_cells:
            sheet_updates[synopsis_sheet] = summary_cells

        # TESTPLAN test result + comment cells
        testplan_sheet = report.mapping["testplan_sheet"]
        cols = report.mapping["testplan_columns"]
        start = report.mapping["testplan_start_row"]
        notes_col = cols.get("engineer_notes")
        header_row = report.mapping.get("testplan_header_row")
        testplan_cells: dict[str, str | None] = {}
        if notes_col and header_row:
            testplan_cells[f"{notes_col}{header_row}"] = "Engineer Notes"
        for offset, row in enumerate(report.rows):
            result_val = _sanitize(row.test_result) or None
            comment_val = _sanitize(row.test_comments) or None
            testplan_cells[f"{cols['test_result']}{start + offset}"] = result_val
            testplan_cells[f"{cols['test_comments']}{start + offset}"] = comment_val
            if notes_col:
                notes_val = _sanitize(row.engineer_notes) or None
                testplan_cells[f"{notes_col}{start + offset}"] = notes_val
        if testplan_cells:
            sheet_updates[testplan_sheet] = testplan_cells

        # ADDITIONAL INFORMATION sections
        additional_sheet = report.mapping.get("additional_sheet")
        add_cells_map = report.mapping.get("additional_cells") or {}
        additional_cells: dict[str, str | None] = {}
        if report.additional_sections and add_cells_map.get("section_1_title") and add_cells_map.get("section_1_body"):
            additional_cells[add_cells_map["section_1_title"]] = report.additional_sections[0].title
            additional_cells[add_cells_map["section_1_body"]] = report.additional_sections[0].body
        if len(report.additional_sections) > 1 and add_cells_map.get("section_2_title") and add_cells_map.get("section_2_body"):
            additional_cells[add_cells_map["section_2_title"]] = report.additional_sections[1].title
            additional_cells[add_cells_map["section_2_body"]] = report.additional_sections[1].body
        if additional_sheet and additional_cells:
            sheet_updates[additional_sheet] = additional_cells

        image_inserts: list[XlsxImageInsert] = []
        issuer_logo = _resolve_electracom_report_logo_path()
        if issuer_logo:
            issuer_width, issuer_height = _scaled_logo_pixels(issuer_logo, width_px=190)
            image_inserts.append(
                XlsxImageInsert(
                    sheet_name=synopsis_sheet,
                    image_path=issuer_logo,
                    cell="A1",
                    width_px=issuer_width,
                    height_px=issuer_height,
                    description="Electracom report logo",
                )
            )
        client_logo = _resolve_uploaded_logo_path(report.branding.logo_path)
        if client_logo:
            client_width, client_height = _scaled_logo_pixels(client_logo, width_px=150)
            image_inserts.append(
                XlsxImageInsert(
                    sheet_name=synopsis_sheet,
                    image_path=client_logo,
                    cell="I1",
                    width_px=client_width,
                    height_px=client_height,
                    description="Client branding logo",
                )
            )

        await asyncio.to_thread(patch_xlsx, report.template_path, path, sheet_updates, image_inserts)
    else:
        # Fallback: no template to preserve — plain openpyxl workbook
        from openpyxl import Workbook

        wb = Workbook()
        wb.active.title = report.summary_section_title
        wb.create_sheet(report.testplan_section_title)
        wb.create_sheet(report.additional_section_title)
        wb.save(str(path))

    return str(path)


async def generate_word_report(
    test_run: Any,
    test_results: list[Any],
    report_config: Any = None,
    include_synopsis: bool = False,
    enabled_test_ids: Optional[list[str]] = None,
    whitelist_entries: Optional[list[Any]] = None,
    template_key: str = "generic",
    branding_settings: Any = None,
    readiness_summary: Optional[dict[str, Any]] = None,
) -> str:
    from docx import Document
    from docx.shared import Inches

    report = build_report_document(
        test_run,
        test_results,
        report_config,
        template_key,
        enabled_test_ids,
        whitelist_entries,
        include_synopsis,
        branding_settings,
        readiness_summary,
    )
    doc = Document()
    issuer_logo = _resolve_electracom_report_logo_path()
    client_logo = _resolve_uploaded_logo_path(report.branding.logo_path)
    header_paragraph = doc.sections[0].header.paragraphs[0]
    if issuer_logo:
        try:
            header_paragraph.add_run().add_picture(str(issuer_logo), width=Inches(1.9))
        except Exception as exc:
            logger.warning("Failed to add Electracom logo to Word report: %s", exc)
    if client_logo:
        try:
            header_paragraph.add_run("  ").add_picture(str(client_logo), width=Inches(1.35))
        except Exception as exc:
            logger.warning("Failed to add client logo to Word report: %s", exc)
    if report.branding.company_name:
        doc.add_paragraph(report.branding.company_name)
    if report.branding.footer_text:
        doc.sections[0].footer.paragraphs[0].text = report.branding.footer_text
    doc.add_heading("IP Device Qualification Report", level=0)
    doc.add_paragraph(f"Template Profile: {TEMPLATE_INFO[report.template_key]['name']}")
    doc.add_heading(report.summary_section_title, level=1)
    table = doc.add_table(rows=0, cols=2)
    for key, label in report.summary_fields:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = str(report.metadata.get(key, ""))
    doc.add_paragraph(f"{report.summary_text_label}: {report.metadata.get('summary_text', '')}")
    doc.add_page_break()
    doc.add_heading(report.testplan_section_title, level=1)
    results_table = doc.add_table(rows=1, cols=len(report.testplan_columns))
    for idx, (_, header) in enumerate(report.testplan_columns):
        results_table.rows[0].cells[idx].text = header
    for item in report.rows:
        row = results_table.add_row()
        for idx, value in enumerate(_report_row_values(item, report.testplan_columns)):
            row.cells[idx].text = value
    doc.add_page_break()
    doc.add_heading(report.additional_section_title, level=1)
    additional_table = doc.add_table(rows=1, cols=2)
    additional_table.rows[0].cells[0].text = "Section"
    additional_table.rows[0].cells[1].text = "Content"
    for section in report.additional_sections:
        row = additional_table.add_row()
        row.cells[0].text = section.title
        row.cells[1].text = section.body
    doc.add_page_break()
    doc.add_heading("Raw Evidence", level=1)
    evidence_table = doc.add_table(rows=1, cols=4)
    for idx, header in enumerate(["Test ID", "Test Name", "Engineer Notes", "Detailed Evidence"]):
        evidence_table.rows[0].cells[idx].text = header
    for item in report.rows:
        row = evidence_table.add_row()
        row.cells[0].text = item.test_number
        row.cells[1].text = item.brief_description
        row.cells[2].text = item.engineer_notes
        row.cells[3].text = item.evidence_detail or item.evidence_summary
    path = _output_path(test_run, template_key, ".docx")
    doc.save(str(path))
    return str(path)


def _dxf_text(value: Any, limit: int = 160) -> str:
    text = _sanitize_text(value)
    text = text.replace("\r\n", " ").replace("\n", " ").replace("\r", " ")
    text = text.replace("\\", "/")
    text = re.sub(r"\s+", " ", text).strip()
    if len(text) > limit:
        text = f"{text[:limit].rstrip()}..."
    return text or " "


def _dxf_text_entity(x: float, y: float, height: float, value: Any, layer: str = "REPORT") -> list[str]:
    return [
        "0", "TEXT",
        "8", layer,
        "10", f"{x:.2f}",
        "20", f"{y:.2f}",
        "30", "0.00",
        "40", f"{height:.2f}",
        "1", _dxf_text(value),
    ]


async def generate_dxf_report(
    test_run: Any,
    test_results: list[Any],
    report_config: Any = None,
    include_synopsis: bool = False,
    enabled_test_ids: Optional[list[str]] = None,
    whitelist_entries: Optional[list[Any]] = None,
    template_key: str = "generic",
    branding_settings: Any = None,
    readiness_summary: Optional[dict[str, Any]] = None,
) -> str:
    """Generate a lightweight DXF CAD interchange report.

    DXF is used intentionally because it is an open CAD exchange format that
    can be generated locally without proprietary DWG tooling.
    """
    report = build_report_document(
        test_run,
        test_results,
        report_config,
        template_key,
        enabled_test_ids,
        whitelist_entries,
        include_synopsis,
        branding_settings,
        readiness_summary,
    )
    path = _output_path(test_run, template_key, ".dxf")

    entities: list[str] = []
    y = 290.0
    entities.extend(_dxf_text_entity(10, y, 4.5, "EDQ IP Device Qualification Report", "TITLE"))
    y -= 8
    entities.extend(_dxf_text_entity(10, y, 2.8, f"Template: {TEMPLATE_INFO[report.template_key]['name']}"))
    y -= 6
    entities.extend(_dxf_text_entity(10, y, 2.8, f"Manufacturer: {report.metadata.get('manufacturer', '')}"))
    y -= 5
    entities.extend(_dxf_text_entity(10, y, 2.8, f"Model: {report.metadata.get('model', '')}"))
    y -= 5
    entities.extend(_dxf_text_entity(10, y, 2.8, f"Overall Result: {report.metadata.get('overall_result', '')}"))
    y -= 8
    entities.extend(_dxf_text_entity(10, y, 2.5, "Test Plan Summary", "TITLE"))
    y -= 5
    entities.extend(_dxf_text_entity(10, y, 2.2, "No. | Brief Description | Result | Comments", "HEADER"))
    y -= 4

    for row in report.rows:
        if y < 12:
            break
        entities.extend(
            _dxf_text_entity(
                10,
                y,
                2.0,
                f"{row.test_number} | {row.brief_description} | {row.test_result or 'PENDING'} | {row.test_comments}",
            )
        )
        y -= 4

    dxf_lines = [
        "0", "SECTION",
        "2", "HEADER",
        "9", "$ACADVER",
        "1", "AC1009",
        "0", "ENDSEC",
        "0", "SECTION",
        "2", "ENTITIES",
        *entities,
        "0", "ENDSEC",
        "0", "EOF",
    ]
    path.write_text("\n".join(dxf_lines) + "\n", encoding="ascii", errors="replace")
    return str(path)


async def _generate_pdf_via_libreoffice(source_path: str) -> str:
    import asyncio
    output_dir_path = Path(settings.REPORT_DIR).resolve()
    source_file = _report_source_file(source_path)
    output_dir = str(output_dir_path)
    proc = await asyncio.create_subprocess_exec(
        "libreoffice", "--headless", "--norestore", "--convert-to", "pdf",
        "--outdir", output_dir, str(source_file),
        stdout=asyncio.subprocess.PIPE,
        stderr=asyncio.subprocess.PIPE,
    )
    _, stderr = await asyncio.wait_for(proc.communicate(), timeout=120)
    if proc.returncode != 0:
        raise RuntimeError(f"LibreOffice conversion failed: {stderr.decode()[:200]}")
    pdf_path = source_file.with_suffix(".pdf")
    if not pdf_path.exists():
        raise RuntimeError(f"Expected PDF file not found at {pdf_path}")
    return str(pdf_path)


async def _generate_pdf_via_fpdf(
    test_run: Any,
    test_results: list[Any],
    report_config: Any = None,
    include_synopsis: bool = False,
    enabled_test_ids: Optional[list[str]] = None,
    whitelist_entries: Optional[list[Any]] = None,
    template_key: str = "generic",
    branding_settings: Any = None,
    readiness_summary: Optional[dict[str, Any]] = None,
) -> str:
    from fpdf import FPDF

    report = build_report_document(
        test_run,
        test_results,
        report_config,
        template_key,
        enabled_test_ids,
        whitelist_entries,
        include_synopsis,
        branding_settings,
        readiness_summary,
    )
    pdf = FPDF()
    if hasattr(pdf, "set_compression"):
        pdf.set_compression(False)
    pdf.set_auto_page_break(auto=True, margin=10)
    pdf.add_page()
    issuer_logo = _resolve_electracom_report_logo_path()
    client_logo = _resolve_uploaded_logo_path(report.branding.logo_path)
    logo_top = pdf.get_y()
    if issuer_logo:
        try:
            pdf.image(str(issuer_logo), x=pdf.l_margin, y=logo_top, w=55)
        except Exception as exc:
            logger.warning("Failed to add Electracom logo to PDF report: %s", exc)
    if client_logo:
        try:
            pdf.image(str(client_logo), x=max(pdf.w - pdf.r_margin - 40, pdf.l_margin + 60), y=logo_top, w=40)
        except Exception as exc:
            logger.warning("Failed to add client logo to PDF report: %s", exc)
    if issuer_logo or client_logo:
        pdf.set_y(logo_top + 14)
    if report.branding.company_name:
        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(0, 8, report.branding.company_name, new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "IP DEVICE QUALIFICATION REPORT", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 6, f"Template Profile: {TEMPLATE_INFO[report.template_key]['name']}", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, report.summary_section_title, new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    for key, label in report.summary_fields:
        _write_pdf_lines(pdf, f"{label}: {str(report.metadata.get(key, ''))}", line_height=6, wrap_width=120)
    pdf.ln(2)
    _write_pdf_lines(pdf, f"{report.summary_text_label}: {report.metadata.get('summary_text', '')}", line_height=5, wrap_width=140)
    pdf.add_page(orientation="L")
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, report.testplan_section_title, new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "B", 8)
    _write_pdf_lines(pdf, " | ".join(label for _, label in report.testplan_columns), line_height=5, wrap_width=180)
    pdf.set_font("Helvetica", "", 8)
    for row in report.rows:
        _write_pdf_lines(
            pdf,
            " | ".join(_report_row_values(row, report.testplan_columns)),
            line_height=5,
            wrap_width=180,
        )
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, report.additional_section_title, new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "B", 10)
    _write_pdf_lines(pdf, "Section | Content", line_height=5, wrap_width=140)
    pdf.set_font("Helvetica", "", 10)
    for section in report.additional_sections:
        _write_pdf_lines(pdf, f"{section.title} | {section.body}", line_height=5, wrap_width=140)
    if report.branding.footer_text:
        _write_pdf_lines(pdf, report.branding.footer_text, line_height=5, wrap_width=140)
    path = _output_path(test_run, template_key, ".pdf")
    pdf.output(str(path))
    return str(path)


async def generate_pdf_report(
    test_run: Any,
    test_results: list[Any],
    report_config: Any = None,
    include_synopsis: bool = False,
    enabled_test_ids: Optional[list[str]] = None,
    whitelist_entries: Optional[list[Any]] = None,
    template_key: str = "generic",
    branding_settings: Any = None,
    readiness_summary: Optional[dict[str, Any]] = None,
) -> str:
    template_key = _validate_template_key(template_key)
    libreoffice_setting = os.getenv("EDQ_USE_LIBREOFFICE", "").strip().lower()
    use_libreoffice = libreoffice_setting in ("1", "true", "yes") or (
        libreoffice_setting not in ("0", "false", "no", "off")
        and shutil.which("libreoffice") is not None
    )

    if use_libreoffice:
        try:
            mapping = _load_mapping(template_key)
            template_path = _TEMPLATES_DIR / TEMPLATE_FILES[template_key]
            if mapping and template_path.exists():
                source_path = await generate_excel_report(
                    test_run,
                    test_results,
                    report_config,
                    template_key,
                    enabled_test_ids,
                    whitelist_entries,
                    include_synopsis,
                    branding_settings,
                    readiness_summary,
                )
            else:
                source_path = await generate_word_report(
                    test_run,
                    test_results,
                    report_config,
                    include_synopsis,
                    enabled_test_ids,
                    whitelist_entries,
                    template_key,
                    branding_settings,
                    readiness_summary,
                )
            return await _generate_pdf_via_libreoffice(source_path)
        except Exception:
            logger.warning("LibreOffice PDF conversion failed, falling back to fpdf2")

    return await _generate_pdf_via_fpdf(
        test_run,
        test_results,
        report_config,
        include_synopsis,
        enabled_test_ids,
        whitelist_entries,
        template_key,
        branding_settings,
        readiness_summary,
    )


async def generate_csv_report(
    test_run: Any,
    test_results: list[Any],
    report_config: Any = None,
    include_synopsis: bool = False,
    enabled_test_ids: Optional[list[str]] = None,
    whitelist_entries: Optional[list[Any]] = None,
    template_key: str = "generic",
    branding_settings: Any = None,
    readiness_summary: Optional[dict[str, Any]] = None,
) -> str:
    report = build_report_document(
        test_run,
        test_results,
        report_config,
        template_key,
        enabled_test_ids,
        whitelist_entries,
        include_synopsis,
        branding_settings,
        readiness_summary,
    )
    path = _output_path(test_run, template_key, ".csv")
    issuer_logo = _resolve_electracom_report_logo_path()
    client_logo = _resolve_uploaded_logo_path(report.branding.logo_path)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow([report.summary_section_title])
        writer.writerow(["Field", "Value"])
        writer.writerow(["Template Profile", TEMPLATE_INFO[report.template_key]["name"]])
        writer.writerow(["Company Name", report.branding.company_name])
        writer.writerow(["Report Logo", _logo_display_name(issuer_logo)])
        if client_logo:
            writer.writerow(["Client Logo", _logo_display_name(client_logo)])
        writer.writerow(
            [
                "Readiness Status",
                f"{report.readiness_summary.get('label', 'Unknown')} ({report.readiness_summary.get('score', 1)}/10)",
            ]
        )
        if report.branding.footer_text:
            writer.writerow(["Footer Text", report.branding.footer_text])
        for key, label in report.summary_fields:
            writer.writerow([label, report.metadata.get(key, "")])
        writer.writerow([report.summary_text_label, report.metadata.get("summary_text", "")])
        writer.writerow([])
        writer.writerow([report.testplan_section_title])
        writer.writerow([label for _, label in report.testplan_columns])
        for row in report.rows:
            writer.writerow(_report_row_values(row, report.testplan_columns))
        writer.writerow([])
        writer.writerow([report.additional_section_title])
        writer.writerow(["Section", "Content"])
        for section in report.additional_sections:
            writer.writerow([section.title, section.body])
    return str(path)
